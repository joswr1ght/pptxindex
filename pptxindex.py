#!/usr/bin/env python

# -*- coding: utf-8 -*-
#

from docx import Document
from itertools import groupby
from operator import itemgetter
from xml.dom.minidom import parse

import sys
import re
import os
import shutil
import glob
import tempfile
import operator
import argparse
import pdb
import string
import zipfile


# Parse the given root recursively (root is intended to be the paragraph element <a:p>
# If we encounter a link-break element a:br, add a new line to global paragraphtext
# If we encounter an element with type TEXT_NODE, append value to paragraphtext
paragraphtext=""
def parse_node(root):
    global paragraphtext
    if root.childNodes:
        for node in root.childNodes:
            if node.nodeType == node.TEXT_NODE:
                paragraphtext += node.nodeValue.encode('ascii', 'ignore')
            if node.nodeType == node.ELEMENT_NODE:
                if node.tagName == 'a:br':
                    paragraphtext += "\n" 
                parse_node(node)

def parseslidecontent(pptxfile, words, booknum, verbose=False):
    global paragraphtext
    skippages = []
    tmpd = tempfile.mkdtemp()
    zipfile.ZipFile(pptxfile).extractall(path=tmpd, pwd=None)

    # Parse slide content first
    path = tmpd + '/ppt/slides/'
    for infile in glob.glob(os.path.join(path, '*.xml')):
        #parse each XML notes file from the notes folder.
        dom = parse(infile)
        noteslist = dom.getElementsByTagName('a:t')
        page = re.sub(r'\D', "", infile.split("/")[-1])
        text = ''

        for node in noteslist:
            xmlTag = node.toxml()
            xmlData = xmlTag.replace('<a:t>', '').replace('</a:t>', '')
            text += " " + xmlData

        # Convert to ascii to simplify
        text = text.encode('ascii', 'ignore')
        if "Course Roadmap" in text:
            if verbose:
                print "Skipping page %d:%s, \"Course Roadmap\" slide."%(booknum,page)
            skippages.append(page)
            words[str(booknum) + ":" + page] = ''
        else:
            words[str(booknum) + ":" + page] = text

    # Next, parse notes content, skipping pages previously identified
    path = tmpd + '/ppt/notesSlides/'
    for infile in glob.glob(os.path.join(path, '*.xml')):

        # Get the slide number
        page = re.sub(r'\D', "", infile.split("/")[-1])

        if page in skippages:
            # Skip this page previously identified with "Course Roadmap" title text
            continue

        # Parse slide notes, adding a space after each paragraph marker, and removing XML markup
        dom = parse(infile)
        paragraphs=dom.getElementsByTagName('a:p')
        for paragraph in paragraphs:
            paragraphtext=""
            parse_node(paragraph)
            #print "DEBUG: " + paragraphtext

            words[str(booknum) + ":" + str(page)] += " " + paragraphtext

    # Remove all the files created with unzip
    shutil.rmtree(tmpd)

    # Remove double-spaces which happens in the content occasionally
    for page in words:
        words[page] = ''.join(ch for ch in words[page] if ch not in set([',']))
        words[page] = re.sub('\. ', " ", words[page])
        words[page] = ' '.join(words[page].split())
    return words


# Validate the contents of the concordance file
def checkconcordance(concordancefile):
    # Declared empty here, just for validating concordance rules
    page = ""
    cspage = ""
    booknum = 0
    pagenum = 0
    wordlist = ""
    cswordlist = ""

    ret=0
    lineno=0
    for line in open(concordancefile,"rU"):
        expression = None
        lineno+=1
        if line[0] == "#" or line == "\n" or line.isspace(): continue
        try:
            key,expression = line.strip().split(";")
        except ValueError:
            # Explicit search term, continue
            continue
        if expression != None:
            try:
                eval(expression)
            except Exception, e:
                ret=1
                sys.stdout.write("Error processing concordance file line " + str(lineno) + ": ")
                sys.stdout.write(str(e))
                sys.stdout.write("\n")
            continue
    return ret

# Take the index of entries, sort and reduce the page numbers
# into ranges (e.g. 1:3,1:4,1:5,1:8 into 1:3-5,1:8
# This is awful and I hope I never have to edit this code.
def indexreduce(index):
    for entry in index:
        matchesbybook = {}
        pages=index[entry]
        for bookpage in pages:
            book,page = bookpage.split(":")
            page = int(page)
            try:
                matchesbybook[book].append(page)
            except KeyError:
                matchesbybook[book] = [page]
        for book in matchesbybook:
            sortedreduced=[]
            matchesbybook[book].sort()
            matchesbybook[book] = numreduce(matchesbybook[book])
            # Return to 1:66, 2:57 format
            
        index[entry] = []
        for book in matchesbybook:
            for page in matchesbybook[book]:
                index[entry].append(book + ":" + page)

    return index

# Take a list of numbers and reduce them into hyphenated ranges for
# sequential values.
def numreduce(data):
    str_list = []
    for k, g in groupby(enumerate(data), lambda (i,x):i-x):
       ilist = map(itemgetter(1), g)
       #print ilist
       if len(ilist) > 1:
          str_list.append('%d-%d' % (ilist[0], ilist[-1]))
       else:
          str_list.append('%d' % ilist[0])
    return str_list 


def indexsort(string):
    book,page = string.split(":")
    page = re.sub('-.*', "", page)
    return int(book)*(int(page)+1000)


def showconcordancehits(index, concordance):
    print "Concordance matches:"
    nohitcount=0
    for key in concordance:
        # The concordance key will not be present in the index list unless it was present
        # in the PPTX file. The except here is for concordance entries that did not produce
        # a match.
        try:
            rangedmatches = len(index[key])
            # left justify the key name with 52 spaces - may need to be adjusted
            print "\t%s%d ranged matches."%(key.ljust(52), rangedmatches)
        except KeyError:
           nohitcount+=1
           print "\t%s0 matches."%(key.ljust(52))

    if nohitcount == 0:
        print "All entries in the concordance file produced matches."
        return

def is_valid_file(parser, arg):
    if not os.path.exists(arg):
        parser.error("The file %s does not exist!" % arg)
    else:
        return open(arg, 'r')  # return an open file handle

if __name__ == "__main__":

    concordancefile = None
    indexoutputfile = None
    testandexit     = None
    templatefile    = None
    verbose         = False

    parser = argparse.ArgumentParser()
    parser.add_argument('-c', '--concordance', dest='concordance', help='concordance filename', type=lambda x: is_valid_file(parser, x), required=True)
    parser.add_argument('-o', '--outfile', dest='outfile', help='MS Word index output filename', metavar='output.docx', type=argparse.FileType('w'))
    parser.add_argument('-i', '--template', dest='template', help='MS Word template file to base index on', metavar='Template.docx', type=argparse.FileType('r'))
    parser.add_argument('-t', '--test', help='Test and validate concordance file syntax, then exit', action='store_true')
    parser.add_argument('-v', '--verbose', help='Verbose output (including 0-hit concordance entries)', action='store_true')
    parser.add_argument('pptxfiles', type=lambda x: is_valid_file(parser, x), action='store', nargs='*')

    args = parser.parse_args()

    # Check all the expressions in the concordance file
    if (checkconcordance(args.concordance.name) != 0):
        sys.stderr.write("Please correct the errors in the concordance file and try again.\n")
        sys.exit(-1)

    # perform a test if need bee
    if args.test:
        print("No errors in the concordance file.")
        sys.exit(0)

    # ensure that pptxfiles are provided
    if len(args.pptxfiles) == 0:
        print("No pptx files provided")
        sys.exit(0)

    if not args.outfile:
        args.outfile = open(args.concordance.name + ".docx", "w")
        
    # Read concordance file and build the dictionary
    concordance = {}
    for line in open(args.concordance.name,"U"):
        if line[0] == "#" or line == "\n" or line.isspace(): continue
        try:
            key,val = line.strip().split(";")
            concordance[key] = val
        except ValueError:
            concordance[line.strip()] = None
    
    args.pptxfiles.sort(key=lambda f: f.name)

    if args.verbose:
        print("Processing PPTX files: %s")%' '.join(os.path.basename(x.name) for x in args.pptxfiles)

    print("Extracting content from PPTX files.")
    wordsbypage = {}
    booknum=1
    for pptxfile in args.pptxfiles:
        if os.path.splitext(pptxfile.name.lower())[1] != ".pptx":
            sys.stderr.write("Cannot process non-pptx filename \"%s\", exiting.\n"%pptxfile)
            sys.exit(-1)
        try:
            # Retrieve slide and notes text for each slide in pptx file
            wordsbypage = parseslidecontent(pptxfile.name, wordsbypage, booknum, args.verbose)
        except zipfile.BadZipfile:
            sys.stderr.write("Invalid pptx file \"%s\", exiting.\n"%pptxfile)
            sys.exit(-1)
        except:
            print "Unexpected error:", sys.exc_info()[0]
            sys.exit(-1)
        booknum+=1

    # Next, iterate through the concordance dictionary, searching for and recording
    # matches for each entry.
    print("Searching for matches with the concordance file.")
    index = {}
    for key in concordance:
        pages = [] # list of page numbers
        for bookpagenum in wordsbypage:
            # To track hits with concordance entries, mark hits for this
            # entry to None by default.

            # These are the variables intended to be accessible by the author in the concordance file
            cspage = wordsbypage[bookpagenum]
            page = wordsbypage[bookpagenum].lower()
            booknum,pagenum = bookpagenum.split(":")
            wordlist = re.split("(?:(?:[^a-zA-Z]+')|(?:'[^a-zA-Z]+))|(?:[^a-zA-Z']+)", page)
            cswordlist = re.split("(?:(?:[^a-zA-Z]+')|(?:'[^a-zA-Z]+))|(?:[^a-zA-Z']+)", cspage)

            # Process the concordance file entry.  If it is None, then use 
            # the key as the search string
            if concordance[key] == None:
                if (key.lower() in page):
                        pages.append(bookpagenum)
            # Else, evaluate the right-side of the concordance entry as a Python expression
            elif eval(concordance[key]):
                pages.append(bookpagenum)

            # If the concordance entry generated some matches, add it to the index list
            if pages != []:
                index[key] = pages

    if args.verbose:
        showconcordancehits(index, concordance)

    # Reduce index entries "1:1,1:2,1:3" to 1:1-3"
    print("Creating index reference ranges.")
    index = indexreduce(index)

    # Sort the reduced index entries numerically
    for page in index:
        index[page] = sorted(index[page], key=indexsort)

    # With index list created, make the Word document
    print("Creating index document.")
    document = Document(args.template)
    #if templatefile != None:
    #    document.add_page_break()
    
    table = document.add_table(rows=0, cols=2, style="Light Shading")
    l2marker=""
    for entry in sorted(index.keys(), key=str.lower):
        if entry == '': continue
        #pdb.set_trace()
        currentmarker = ord(entry[0].upper())
        if currentmarker > 64: # "A" or after
            if l2marker != currentmarker:
                document.add_heading(entry[0].upper(), level=1)
                table = document.add_table(rows=0, cols=2, style="Light Shading")
                l2marker=currentmarker
        row_cells = table.add_row().cells
        row_cells[0].text = entry
        row_cells[1].text = ", ".join(index[entry])

    args.outfile.close()
    document.save(args.outfile.name)
    print("Done.")
